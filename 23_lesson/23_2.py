from docx import Document


def markdown_to_docx(t):
    doc = Document()
    t = open(t, 'r', encoding='utf8')
    l = t.readlines()
    t = ''
    for i in l:
        t += str(i)
    l = t.split('\n')
    doc.add_heading(l[0], level=0)
    for l in l[1:]:
        if l:
            i = 0
            f = 0
            while i < len(l):
                if l[i] == '#' and f == 1:
                    if i > 1 and l[i - 1] == "`":
                        pass
                    else:
                        l = l[:i] + l[i + 1:]
                elif l[i] != '#':
                    f = 1
                i += 1
            f2 = 1
            if l[0] == '#':
                if l.count('#') == 1:
                    doc.add_heading(l[2:], level=1)
                elif l.count('#') == 2:
                    doc.add_heading(l[3:], level=2)
                elif l.count('#') == 3:
                    doc.add_heading(l[4:], level=3)
                elif l.count('#') == 4:
                    doc.add_heading(l[5:], level=4)
                elif l.count('#') == 5:
                    doc.add_heading(l[6:], level=5)
                elif l[:7].count('#') == 6:
                    doc.add_heading(l[7:], level=6)
            elif str(l[:2]) == '- ':
                doc.add_paragraph(l[2:], style='List Bullet')
            elif str(l[:2]) == '* ':
                doc.add_paragraph(l[2:], style='List Bullet')
            elif str(l[:2]) == '+ ':
                doc.add_paragraph(l[2:], style='List Bullet')
            elif l[0].isdigit() and l[1] == '.':
                doc.add_paragraph(l[3:], style='List Number')
            elif l.count('_') == 12:
                r = doc.add_paragraph()
                r.add_run(l[2:22]).bold = True
                k = r.add_run(l[23:29] + ' ')
                k.bold = True
                k.italic = True
                f2 = 0
                r.add_run(l[31:43]).bold = True
                r.add_run(l[45:60])
                r.add_run(l[61:63]).italic = True
                k = r.add_run(l[65:79] + ' ')
                k.bold = True
                k.italic = True
                r.add_run(l[81:-1]).italic = True
            elif l[:3].count('_') == 1 or l[:3].count('*') == 1 and f2 == 1:
                doc.add_paragraph().add_run(l[1:-1]).italic = True
            elif l[:3].count('_') == 2 or l[:3].count('*') == 2 and f2 == 1:
                doc.add_paragraph().add_run(l[2:-2]).bold = True
            elif l[:3].count('_') == 3 or l[:3].count('*') == 3 and f2 == 1:
                r = doc.add_paragraph().add_run(l[3:-3])
                r.bold = True
                r.italic = True
            else:
                doc.add_paragraph(l)
        else:
            doc.add_paragraph()
    doc.save('t3.docx')


markdown_to_docx('03_in.txt')